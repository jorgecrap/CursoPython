
import openpyxl
import docx

# Excels

'''
excel_document = openpyxl.load_workbook('test.xlsx')
hojas = excel_document.get_sheet_names()
clientes = excel_document.get_sheet_by_name('Clientes')

celda = clientes.cell(row = 3, column = 2).value

rango = clientes['A1:B15']

for linea in rango:
	for celda in linea:
		print (celda.value)

'''


# Words

doc = docx.Document('demo.docx')
doc2 = docx.Document('demo2.docx')

tb = doc.tables[0]
tb2 = doc.tables[0]

doc.save('demo.docx')
doc2.save('demo2.docx')

